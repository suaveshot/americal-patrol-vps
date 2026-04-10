# sales_pipeline/proposal/create_template.py
"""
Builds proposal_template.docx — Americal Patrol proposal template.

All variable fields use {{PLACEHOLDER}} tokens that proposal_generator.py
will replace at send time.

Run:
    python -m sales_pipeline.proposal.create_template
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = BASE_DIR / "proposal_template.docx"

FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = Pt(12)
FONT_SIZE_HEADER = Pt(12)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_font(run, bold=False, size=None, color=None):
    run.font.name = FONT_NAME
    run.font.bold = bold
    run.font.size = size or FONT_SIZE_BODY
    if color:
        run.font.color.rgb = RGBColor(*color)


def _para(doc, text="", bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
          size=None, space_before=None, space_after=None):
    """Add a paragraph with consistent font settings."""
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    if space_before is not None:
        fmt.space_before = space_before
    if space_after is not None:
        fmt.space_after = space_after
    if text:
        run = p.add_run(text)
        _set_font(run, bold=bold, size=size)
    return p


def _inline(para, text, bold=False, size=None):
    """Add a run to an existing paragraph."""
    run = para.add_run(text)
    _set_font(run, bold=bold, size=size)
    return run


def _set_margins(section, top=1.0, bottom=1.0, left=1.25, right=1.25):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def _add_horizontal_rule(doc):
    """Add a paragraph with a bottom border to simulate a horizontal rule."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    return p


def _no_space(para):
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)


# ---------------------------------------------------------------------------
# Page 1 — Cover Letter
# ---------------------------------------------------------------------------

def build_cover_letter(doc):
    # --- Letterhead header ---
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _no_space(p)
    _inline(p, "AMERICAL PATROL, INC.", bold=True, size=Pt(14))

    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _no_space(p)
    _inline(p, "CALIFORNIA LICENSE PPO 9557", bold=False, size=Pt(11))

    _add_horizontal_rule(doc)

    # Two-column address/phone block — simulated with a table (no borders)
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = "Table Grid"
    # Remove all borders
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border_el = OxmlElement(f"w:{side}")
                border_el.set(qn("w:val"), "none")
                tcBorders.append(border_el)
            tcPr.append(tcBorders)

    def _cell_run(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=None):
        p = cell.paragraphs[0]
        p.alignment = align
        _no_space(p)
        run = p.add_run(text)
        _set_font(run, bold=bold, size=size or Pt(11))

    _cell_run(tbl.cell(0, 0), "3301 Harbor Blvd., Oxnard, CA 93035")
    _cell_run(tbl.cell(0, 1), "LA & OC  (714) 521-0855", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_run(tbl.cell(1, 0), "www.americalpatrol.com")
    _cell_run(tbl.cell(1, 1), "Ventura  (805) 844-9433", align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_paragraph()  # breathing room

    # --- Date ---
    p = _para(doc, "{{DATE}}")
    _no_space(p)

    doc.add_paragraph()

    # --- Recipient address block ---
    for token in ("{{RECIPIENT_NAME}}", "{{COMPANY_NAME}}", "{{STREET_ADDRESS}}", "{{CITY_STATE_ZIP}}"):
        p = _para(doc, token)
        _no_space(p)

    doc.add_paragraph()

    # --- RE line ---
    p = _para(doc)
    _no_space(p)
    _inline(p, "RE: ", bold=True)
    _inline(p, "Security Services")

    doc.add_paragraph()

    # --- Salutation ---
    _para(doc, "Dear {{FIRST_NAME}},")

    doc.add_paragraph()

    # --- Opening line ---
    _para(
        doc,
        "Thank you for allowing Americal Patrol, Inc., the opportunity to bid "
        "for your security services.",
        space_after=Pt(6),
    )

    # --- Boilerplate paragraphs ---
    boilerplate = [
        (
            "Americal Patrol, Inc. is licensed by the Bureau of Security and "
            "Investigative Services and has been conducting business in California "
            "since 1986. We are a medium size company that offers a full line of "
            "services, such as: armed and unarmed uniform guards, and private patrol "
            "service."
        ),
        (
            "We maintain General Liability Insurance in the aggregate amount of "
            "$3,000,000. We also maintain workers compensation with State Fund in the "
            "amount of $1,000,000. Certificates of Insurance will be issued upon "
            "request."
        ),
        (
            "I realize that in today\u2019s climate everyone is budget conscious. "
            "However, keep in mind that security is a mindset, not a specific product. "
            "Security is also a continuous process, not a short-term project. More "
            "importantly, security is a piece of the risk management function. We are "
            "a team player, and we look forward to joining your risk management team."
        ),
        (
            "I am confident that Americal Patrol, Inc. will meet or exceed your "
            "expectations of a contract guard company."
        ),
        (
            "Should you have any questions or need additional information please give "
            "me a call."
        ),
    ]

    for text in boilerplate:
        _para(doc, text, space_after=Pt(6))

    doc.add_paragraph()

    # --- Closing ---
    _para(doc, "Respectfully,")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    for line in (
        "Sam A. Alarcon",
        "Sam A. Alarcon, Vice President",
        "Americal Patrol, Inc.",
        "SAlarcon@americalpatrol.com",
    ):
        p = _para(doc, line)
        _no_space(p)


# ---------------------------------------------------------------------------
# Page 2 — Attachment A
# ---------------------------------------------------------------------------

def build_attachment_a(doc):
    # Page break
    doc.add_page_break()

    # --- Title ---
    p = _para(doc, 'ATTACHMENT "A"', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14))
    p.paragraph_format.space_after = Pt(6)

    # --- Company name + subtitle ---
    p = _para(doc, "{{COMPANY_NAME}}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _no_space(p)
    p = _para(doc, "SECURITY COVERAGE AND PROPOSAL", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(12)

    _add_horizontal_rule(doc)

    # --- Info table (label: value) ---
    # We use a borderless 2-column table to get clean label/value alignment
    info_rows = [
        ("Date:",         "{{DATE}}"),
        ("Prepared for:", "{{RECIPIENT_NAME}}, {{RECIPIENT_TITLE}}"),
        ("",              "{{COMPANY_NAME}}"),
        ("Prepared by:",  "Sam Alarcon, Vice President"),
        ("",              "Americal Patrol, Inc."),
        ("Location:",     "{{LOCATION_ADDRESS}}"),
    ]

    tbl = doc.add_table(rows=len(info_rows), cols=2)
    tbl.style = "Table Grid"

    # Set column widths: label ~1.5 in, value ~4.5 in
    for i, row in enumerate(tbl.rows):
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.5)

        label, value = info_rows[i]

        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border_el = OxmlElement(f"w:{side}")
                border_el.set(qn("w:val"), "none")
                tcBorders.append(border_el)
            tcPr.append(tcBorders)

        p0 = row.cells[0].paragraphs[0]
        _no_space(p0)
        r0 = p0.add_run(label)
        _set_font(r0, bold=bool(label))

        p1 = row.cells[1].paragraphs[0]
        _no_space(p1)
        r1 = p1.add_run(value)
        _set_font(r1)

    doc.add_paragraph()

    # --- Description ---
    p = _para(doc)
    _no_space(p)
    _inline(p, "Description:  ", bold=True)
    _inline(p, "{{SERVICE_DESCRIPTION}}")
    doc.add_paragraph()

    # --- Hourly Rates ---
    p = _para(doc, "Hourly Rates:", bold=True)
    p.paragraph_format.space_after = Pt(2)

    p = _para(doc)
    _no_space(p)
    _inline(p, "{{RATE_LABEL}}", bold=True)
    _inline(p, ":  {{RATE_AMOUNT}}")

    doc.add_paragraph()

    # --- Billing note ---
    note_text = (
        "Note: Americal Patrol will bill for the time of any officer who must appear "
        "for a deposition, testimony or attend a hearing or court case related to the "
        "service at the contracted property. Said compensation will be at the overtime "
        "rate."
    )
    p = _para(doc, note_text, size=Pt(10))
    p.paragraph_format.space_after = Pt(8)

    _add_horizontal_rule(doc)

    # --- Monthly Cost ---
    p = _para(doc)
    _inline(p, "Monthly Cost:  ", bold=True, size=Pt(13))
    _inline(p, "{{MONTHLY_COST}}", bold=True, size=Pt(13))
    p.paragraph_format.space_after = Pt(2)

    p = _para(doc)
    _no_space(p)
    _inline(p, "({{COST_BASIS}})")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def create_template(output_path: Path = OUTPUT_PATH) -> Path:
    doc = Document()

    # Page setup
    section = doc.sections[0]
    _set_margins(section)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE_BODY

    build_cover_letter(doc)
    build_attachment_a(doc)

    doc.save(output_path)
    print(f"Template saved: {output_path}")
    return output_path


if __name__ == "__main__":
    create_template()
