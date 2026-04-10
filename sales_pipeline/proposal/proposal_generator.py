"""
Sales Pipeline — Proposal Generator (GHL Estimates + Branded PDF)
Creates GHL estimates with Claude-enhanced service descriptions
and generates a branded PDF proposal to attach.
"""

import logging
import re
from dataclasses import dataclass
from datetime import date, timedelta
from pathlib import Path

import anthropic
from docx import Document

from sales_pipeline.config import (
    ANTHROPIC_API_KEY,
    GHL_LOCATION_ID,
    GHL_USER_ID,
    PIPELINE_ID,
    PROPOSAL_SENT_STAGE,
    GMAIL_SENDER,
    TEMPLATE_FILE,
)

log = logging.getLogger(__name__)

GENERATED_DIR = Path(__file__).resolve().parent / "generated"

# Americal Patrol business details (constant)
BUSINESS_DETAILS = {
    "name": "Americal Patrol",
    "phone": "(805) 515-3834",
    "website": "americalpatrol.com",
    "address": {
        "addressLine1": "Oxnard, CA",
    },
}

# Standard terms & conditions (HTML)
TERMS_AND_CONDITIONS = """<p><strong>Terms &amp; Conditions:</strong></p>
<ul>
<li>Services begin upon signed agreement and receipt of first payment.</li>
<li>30-day written cancellation notice required.</li>
<li>Americal Patrol carries full liability insurance and workers' compensation coverage.</li>
<li>All patrol officers are licensed, uniformed, and background-checked.</li>
<li>Monthly invoicing — payment due within 15 days of invoice date.</li>
<li>Service scope adjustments available with 7-day advance notice.</li>
</ul>"""

# Property-type-specific duties for the service description
_PROPERTY_DUTIES = {
    "hoa": (
        "patrolling the property, ensuring all gates & doors are locked and secured, "
        "patrol communal areas such as pools, clubhouse, laundry rooms, etc., "
        "citing and towing vehicles, and challenging suspicious people/transients"
    ),
    "commercial": (
        "patrolling the property, ensuring all gates & doors are locked and secured, "
        "checking loading docks and storage areas, monitoring parking structures, "
        "and challenging suspicious people/transients"
    ),
    "industrial": (
        "patrolling the property perimeter, ensuring all gates & doors are locked and secured, "
        "monitoring cargo and equipment areas, verifying truck access, "
        "and challenging suspicious people/transients"
    ),
    "other": (
        "patrolling the property, ensuring all gates & doors are locked and secured, "
        "and challenging suspicious people/transients"
    ),
}

_SERVICE_DESC_SYSTEM = (
    "You are writing service description copy for Americal Patrol, "
    "a licensed security patrol company in Ventura County, CA. "
    "Match this exact tone and structure: direct, professional, no filler. "
    "Write in third person as if describing the service in a formal proposal."
)


WEEKS_PER_MONTH = 4.33  # average weeks in a month

# ── Service Presets ─────────────────────────────────────────────────────────
# Each preset defines the line-item details for a standard service package.
# Guard presets require weekly_hours to calculate monthly cost.
SERVICE_PRESETS = {
    "patrol-1x": {
        "service_name": "Mobile Vehicle Patrol — 1x Daily",
        "patrol_frequency": "1 daily vehicle patrol",
        "unit_amount": 500.00,
        "quantity": 1,
        "patrols_per_day": 1,
        "description_hint": "1 patrol per day",
    },
    "patrol-2x": {
        "service_name": "Mobile Vehicle Patrol — 2x Daily",
        "patrol_frequency": "2 daily vehicle patrols",
        "unit_amount": 1000.00,
        "quantity": 1,
        "patrols_per_day": 2,
        "description_hint": "2 patrols per day",
    },
    "patrol-3x": {
        "service_name": "Mobile Vehicle Patrol — 3x Daily",
        "patrol_frequency": "3 daily vehicle patrols",
        "unit_amount": 1500.00,
        "quantity": 1,
        "patrols_per_day": 3,
        "description_hint": "3 patrols per day",
    },
    "guard-unarmed": {
        "service_name": "Unarmed Standing Guard",
        "patrol_frequency": "dedicated unarmed standing guard",
        "hourly_rate": 28.00,
        "description_hint": "unarmed guard service",
    },
    "guard-armed": {
        "service_name": "Armed Standing Guard",
        "patrol_frequency": "dedicated armed standing guard",
        "hourly_rate": 31.50,
        "description_hint": "armed guard service",
    },
}


@dataclass
class EstimateInput:
    """Input for creating a GHL estimate."""
    contact_id: str
    property_type: str  # hoa, commercial, industrial, other
    patrol_frequency: str  # e.g. "2 daily vehicle patrols"
    patrol_hours: str  # e.g. "after 11pm, 1 hour each patrol"
    service_name: str  # line item name, e.g. "Vehicle Patrol Service"
    unit_amount: float  # monthly price (or calculated from hourly × hours)
    quantity: int  # line item qty (typically 1 for monthly pricing)
    opportunity_id: str = ""  # optional — link to existing opportunity
    patrols_per_day: int = 0  # patrol services: used for per-check rate calc


def resolve_preset(data: dict) -> dict:
    """
    Expand a template-based input dict into full EstimateInput fields.

    Accepts either:
      {"contact_id": "...", "template": "patrol-2x"}
      {"contact_id": "...", "template": "guard-unarmed", "weekly_hours": 40}
      {"contact_id": "...", "service_name": "...", ...}  (custom — passed through)

    Returns a dict ready for EstimateInput(**result).
    """
    template_name = data.get("template")
    if not template_name:
        # No template — treat as custom input, pass through
        return data

    if template_name not in SERVICE_PRESETS:
        raise ValueError(
            f"Unknown template '{template_name}'. "
            f"Available: {', '.join(SERVICE_PRESETS.keys())}"
        )

    preset = SERVICE_PRESETS[template_name]
    result = {
        "contact_id": data["contact_id"],
        "property_type": data.get("property_type", "commercial"),
        "patrol_hours": data.get("patrol_hours", "9pm to 5am"),
        "patrol_frequency": preset["patrol_frequency"],
        "service_name": preset["service_name"],
        "opportunity_id": data.get("opportunity_id", ""),
    }

    if "hourly_rate" in preset:
        # Guard preset — requires weekly_hours
        weekly_hours = data.get("weekly_hours")
        if not weekly_hours:
            raise ValueError(
                f"Template '{template_name}' requires 'weekly_hours' field"
            )
        weekly_hours = float(weekly_hours)
        monthly_hours = round(weekly_hours * WEEKS_PER_MONTH, 1)
        monthly_cost = round(preset["hourly_rate"] * monthly_hours, 2)
        result["unit_amount"] = monthly_cost
        result["quantity"] = 1
        result["service_name"] = (
            f"{preset['service_name']} — {int(weekly_hours)} hrs/week"
        )
        result["patrol_frequency"] = (
            f"{preset['patrol_frequency']}, {int(weekly_hours)} hours per week"
        )
        result["patrol_hours"] = data.get("patrol_hours", "as scheduled")
    else:
        # Patrol preset — flat monthly pricing
        result["unit_amount"] = preset["unit_amount"]
        result["quantity"] = preset["quantity"]
        result["patrols_per_day"] = preset.get("patrols_per_day", 0)

    return result


def _call_claude(prompt: str) -> str:
    """Call Claude API and return the text response."""
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY())
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=512,
        system=_SERVICE_DESC_SYSTEM,
        messages=[{"role": "user", "content": prompt}],
    )
    return response.content[0].text.strip()


def build_service_description(inp: EstimateInput) -> str:
    """
    Call Claude to generate a 1-2 paragraph service description.
    """
    duties = _PROPERTY_DUTIES.get(inp.property_type.lower(), _PROPERTY_DUTIES["other"])
    prompt = (
        f"Write 1-2 short paragraphs describing the security patrol service for a proposal. "
        f"Use this exact structure and style:\n\n"
        f"'Americal Patrol will perform [frequency] between the times of [hours]. "
        f"Officers are a visible deterrent to crime, and their duties include but are not limited to: "
        f"[duties relevant to property type].'\n\n"
        f"Fill in with:\n"
        f"- Frequency: {inp.patrol_frequency}\n"
        f"- Hours: {inp.patrol_hours}\n"
        f"- Property type: {inp.property_type}\n"
        f"- Duties: {duties}\n\n"
        f"Keep it concise, professional, no filler. Do not add greetings or sign-offs."
    )
    return _call_claude(prompt)


def build_estimate_body(inp: EstimateInput, contact: dict, service_description: str) -> dict:
    """
    Build the GHL estimate request body.

    Args:
        inp: EstimateInput with service/pricing details
        contact: GHL contact dict (from ghl_client.get_contact)
        service_description: Claude-generated description

    Returns:
        Dict ready to POST to /invoices/estimate
    """
    location_id = GHL_LOCATION_ID()
    today = date.today()
    expiry = today + timedelta(days=30)

    company_name = contact.get("companyName", "")
    first_name = contact.get("firstName", "")
    last_name = contact.get("lastName", "")
    full_name = f"{first_name} {last_name}".strip() or contact.get("name", "")
    email = contact.get("email", "")
    phone = contact.get("phone", "")

    # Build contact address from GHL contact fields
    address = {}
    if contact.get("address1"):
        address["addressLine1"] = contact["address1"]
    if contact.get("city"):
        address["city"] = contact["city"]
    if contact.get("state"):
        address["state"] = contact["state"]
    if contact.get("postalCode"):
        address["postalCode"] = contact["postalCode"]
    if contact.get("country"):
        address["country"] = contact["country"]

    label = company_name or full_name
    estimate_name = f"Patrol Proposal — {label}"[:40]

    body = {
        "altId": location_id,
        "altType": "location",
        "name": estimate_name,
        "title": "ESTIMATE",
        "issueDate": today.isoformat(),
        "expiryDate": expiry.isoformat(),
        "currency": "USD",
        "businessDetails": BUSINESS_DETAILS,
        "contactDetails": {
            "id": inp.contact_id,
            "name": full_name,
            "email": email,
            "phoneNo": phone,
            "companyName": company_name,
        },
        "items": [
            {
                "name": inp.service_name,
                "description": service_description,
                "currency": "USD",
                "amount": inp.unit_amount,
                "qty": inp.quantity,
                "type": "one_time",
            }
        ],
        "discount": {"value": 0, "type": "percentage"},
        "frequencySettings": {"enabled": False},
        "termsNotes": TERMS_AND_CONDITIONS,
        "liveMode": True,
    }

    if address:
        body["contactDetails"]["address"] = address

    user_id = GHL_USER_ID()
    if user_id:
        body["userId"] = user_id

    return body


# ── Branded PDF Generation ──────────────────────────────────────────────────

def _replace_in_paragraph(paragraph, replacements: dict) -> None:
    """Replace placeholder tokens in a paragraph, handling split runs."""
    full_text = "".join(run.text for run in paragraph.runs)
    if "{{" not in full_text:
        return
    new_text = full_text
    for placeholder, value in replacements.items():
        new_text = new_text.replace(placeholder, value)
    if new_text == full_text:
        return
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def _fill_template(template_path: Path, output_path: Path, replacements: dict) -> None:
    """Fill .docx template placeholders in paragraphs and table cells."""
    doc = Document(str(template_path))
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _format_date(d: date = None) -> str:
    """Format a date as 'March 6th, 2026'."""
    if d is None:
        d = date.today()
    day = d.day
    suffix = (
        "th" if 11 <= day <= 13
        else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    )
    return d.strftime(f"%B {day}{suffix}, %Y")


def _safe_filename(name: str) -> str:
    """Convert a name to a safe filename stem."""
    name = re.sub(r"[^\w\s-]", "", name)
    name = re.sub(r"\s+", "_", name.strip())
    return name


def generate_branded_pdf(
    inp: EstimateInput,
    contact: dict,
    service_description: str,
) -> Path:
    """
    Generate the branded .docx/.pdf proposal matching Sam's format.

    Uses the existing proposal_template.docx with all 13+ placeholders.
    Returns the output file path (.pdf if conversion works, .docx otherwise).
    """
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)

    first_name = contact.get("firstName", "")
    last_name = contact.get("lastName", "")
    full_name = f"{first_name} {last_name}".strip()
    company_name = contact.get("companyName", "")
    title = contact.get("title", contact.get("customField", {}).get("title", ""))

    # Build address parts from GHL contact
    address1 = contact.get("address1", "")
    city = contact.get("city", "")
    state = contact.get("state", "")
    postal = contact.get("postalCode", "")
    city_state_zip = f"{city}, {state} {postal}".strip(", ")

    # Cost display
    monthly_cost = f"${inp.unit_amount * inp.quantity:,.2f}"

    # Build rate line and cost basis based on service type
    if inp.patrols_per_day > 0:
        # Patrol: show transparent per-check pricing
        yearly = inp.unit_amount * inp.quantity * 12
        per_check = round(yearly / 365 / inp.patrols_per_day, 2)
        rate_line = f"Vehicle Patrol Rate: ${per_check:,.2f} per check"
        n = inp.patrols_per_day
        cost_basis = (
            f"(Based on {n} patrol{'s' if n > 1 else ''} per day "
            f"@ ${per_check:,.2f} per check)"
        )
    else:
        rate_line = f"{inp.service_name}: ${inp.unit_amount:,.2f}/month"
        cost_basis = f"(Based on {inp.patrol_frequency})"

    replacements = {
        "{{RECIPIENT_NAME}}":       full_name,
        "{{FIRST_NAME}}":           first_name or full_name,
        "{{RECIPIENT_TITLE}}":      title or "Property Manager",
        "{{COMPANY_NAME}}":         company_name or full_name,
        "{{COMPANY_NAME_UPPER}}":   (company_name or full_name).upper(),
        "{{STREET_ADDRESS}}":       address1,
        "{{CITY_STATE_ZIP}}":       city_state_zip,
        "{{LOCATION_ADDRESS}}":     f"{address1}, {city_state_zip}".strip(", "),
        "{{DATE}}":                 _format_date(),
        "{{SERVICE_DESCRIPTION}}":  service_description,
        "{{RATE_LINE}}":            rate_line,
        "{{MONTHLY_COST}}":         monthly_cost,
        "{{COST_BASIS}}":           cost_basis,
    }

    stem = _safe_filename(company_name or full_name) + "_Security_Proposal"
    docx_path = GENERATED_DIR / f"{stem}.docx"
    _fill_template(TEMPLATE_FILE, docx_path, replacements)
    log.info("Generated branded proposal: %s", docx_path)

    # Attempt PDF conversion
    pdf_path = GENERATED_DIR / f"{stem}.pdf"
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        log.info("PDF conversion successful: %s", pdf_path)
        return pdf_path
    except Exception as e:
        log.warning("PDF conversion skipped (%s) — using .docx", e)
        return docx_path


def create_and_send_estimate(ghl_client, inp: EstimateInput, send: bool = True) -> dict:
    """
    Full pipeline:
    1. Fetch contact from GHL
    2. Generate Claude service description
    3. Build estimate body
    4. Create estimate in GHL
    5. Optionally send it
    6. Update opportunity stage + monetary value

    Returns dict with estimate_id, total, status, and contact info.
    """
    # Step 1: Fetch contact
    contact = ghl_client.get_contact(inp.contact_id)
    log.info("Fetched contact: %s", contact.get("firstName", inp.contact_id))

    # Step 2: Generate service description
    service_description = build_service_description(inp)
    log.info("Generated service description (%d chars)", len(service_description))

    # Step 3: Generate branded PDF proposal
    pdf_path = generate_branded_pdf(inp, contact, service_description)
    log.info("Branded proposal: %s", pdf_path)

    # Step 4: Build estimate body
    body = build_estimate_body(inp, contact, service_description)

    # Step 5: Create estimate
    result = ghl_client.create_estimate(body)
    estimate_id = result.get("_id", result.get("id", ""))
    total = result.get("total", inp.unit_amount * inp.quantity)
    log.info("Created estimate %s (total: $%.2f)", estimate_id, total)

    # Step 6: Send estimate
    if send and estimate_id:
        user_id = GHL_USER_ID()
        from_email = GMAIL_SENDER()
        sent_from = {"fromName": "Sam Alarcon", "fromEmail": from_email} if from_email else None

        try:
            ghl_client.send_estimate(
                estimate_id,
                action="sms_and_email",
                user_id=user_id,
                sent_from=sent_from,
            )
            log.info("Estimate %s sent to contact", estimate_id)
        except Exception as e:
            log.error("Failed to send estimate %s: %s", estimate_id, e)

    # Step 7: Update opportunity if linked
    if inp.opportunity_id:
        try:
            proposal_stage = PROPOSAL_SENT_STAGE()
            updates = {"monetaryValue": total}
            if proposal_stage:
                updates["pipelineStageId"] = proposal_stage
            ghl_client.update_opportunity(inp.opportunity_id, updates)
            log.info("Updated opportunity %s (stage + value)", inp.opportunity_id)
        except Exception as e:
            log.error("Failed to update opportunity: %s", e)

    return {
        "estimate_id": estimate_id,
        "total": total,
        "contact_name": body["contactDetails"]["name"],
        "company_name": body["contactDetails"].get("companyName", ""),
        "status": "sent" if send else "draft",
        "proposal_pdf": str(pdf_path),
    }
